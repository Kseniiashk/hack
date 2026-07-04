"""
Ingest layer — превращает разнородные входные файлы кейса в единую структуру TailingsReport.

Главный принцип: НЕ привязываемся к номерам строк. Ищем якорные кириллические
подписи (метки классов крупности и минералов) и читаем блок под ними. Это
переживает разную длину файлов (Пример 4 длиннее), сдвиги и артефакты #REF!.
"""
from __future__ import annotations
from dataclasses import dataclass, field, asdict
from typing import Optional
import os
import re
import json

try:
    import openpyxl
except ImportError:  # pragma: no cover
    openpyxl = None

# --- Доменные константы -------------------------------------------------------

# Анонимизация организаторов: Элемент 28 = Ni, Элемент 29 = Cu.
ELEMENTS = {
    "28": {"symbol": "Ni", "name": "Никель"},
    "29": {"symbol": "Cu", "name": "Медь"},
}

# Классы крупности в порядке от крупного к мелкому (мкм).
SIZE_CLASSES = ["+125", "+71", "-71 + 45", "-45 + 20", "-20 + 10", "-10"]

# Минералы и их извлекаемость текущей флотацией.
# liberated  — раскрытое зерно, флотируется как есть.
# locked     — сросток, заперт в породе -> нужно доизмельчение, чтобы извлечь.
# tail       — принципиально не извлекается текущей технологией (в породе/решётке).
MINERAL_KIND = {
    "Раскрытый Pnt/Cp": "liberated",
    "Закрытый Pnt/Cp": "locked",
    "Миллерит": "liberated",          # для Ni извлекаемый
    "Примесь в пирротине": "tail",
    "Силикатная форма/Валлериит": "tail",
    "Пирит/Другие Элемент 29": "tail",
    "Потери (расписать)": "tail",
    "Свободный слот": "tail",
}


# --- Структуры данных ---------------------------------------------------------

@dataclass
class MineralSplit:
    """Минералогия одного элемента внутри одного класса крупности."""
    minerals: dict = field(default_factory=dict)   # name -> tons
    recoverable_t: float = 0.0
    unrecoverable_t: float = 0.0

    def liberated_t(self) -> float:
        return sum(t for m, t in self.minerals.items()
                   if MINERAL_KIND.get(m) == "liberated")

    def locked_t(self) -> float:
        return sum(t for m, t in self.minerals.items()
                   if MINERAL_KIND.get(m) == "locked")


@dataclass
class SizeClassRow:
    """Строка распределения металла по классу крупности."""
    size: str
    mass_share_pct: float = 0.0        # доля класса по массе, %
    el28_share_pct: float = 0.0        # доля потерь Ni в этом классе, %
    el28_t: float = 0.0                # т Ni в классе
    el29_share_pct: float = 0.0
    el29_t: float = 0.0
    # минералогия по классу
    min28: MineralSplit = field(default_factory=MineralSplit)
    min29: MineralSplit = field(default_factory=MineralSplit)


@dataclass
class TailingsReport:
    source_name: str = ""
    plant: str = ""
    feed_mass_t: float = 0.0           # шихта, СМТ
    tail_mass_t: float = 0.0           # масса отвальных хвостов, СМТ
    tail_el28_pct: float = 0.0         # содержание Ni в хвостах, %
    tail_el28_t: float = 0.0           # т Ni в хвостах (потери)
    tail_el29_pct: float = 0.0
    tail_el29_t: float = 0.0
    classes: list = field(default_factory=list)   # list[SizeClassRow]
    warnings: list = field(default_factory=list)

    def to_json(self) -> str:
        return json.dumps(asdict(self), ensure_ascii=False, indent=2)


# --- Утилиты ------------------------------------------------------------------

def _num(v) -> float:
    """Безопасное приведение к числу. #REF!, None, текст -> 0.0."""
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(",", ".")
    if not s or s.startswith("#") or s.upper() in ("N/A", "NAN"):
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def _norm(s) -> str:
    """Нормализация подписи: схлопываем пробелы, убираем хвостовые запятые/двоеточия."""
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()


def _canon_mineral(label: str) -> Optional[str]:
    """Сопоставляет подпись минерала из ячейки с каноническим именем каталога."""
    l = _norm(label).lower()
    if not l:
        return None
    if l.startswith("раскрыт"):
        return "Раскрытый Pnt/Cp"
    if l.startswith("закрыт"):
        return "Закрытый Pnt/Cp"
    if l.startswith("миллерит"):
        return "Миллерит"
    if l.startswith("примесь в пирротине"):
        return "Примесь в пирротине"
    if l.startswith("силикатн"):
        return "Силикатная форма/Валлериит"
    if l.startswith("пирит"):
        return "Пирит/Другие Элемент 29"
    if l.startswith("потери"):
        return "Потери (расписать)"
    if l.startswith("свободный слот"):
        return "Свободный слот"
    return None


def _canon_size(label: str) -> Optional[str]:
    """Определяет класс крупности по подписи ячейки (заголовок блока)."""
    l = _norm(label).lower().replace("мкм", "").strip()
    l = l.replace(" ", "")
    mapping = {
        "+125": "+125",
        "+71": "+71",
        "-71+45": "-71 + 45",
        "-45+20": "-45 + 20",
        "-20+10": "-20 + 10",
        "-10": "-10",
    }
    return mapping.get(l)


# --- Основной парсер xlsx -----------------------------------------------------

def parse_tailings_xlsx(path: str, plant: str = "") -> TailingsReport:
    """Парсит файл 'Хвосты …xlsx' в TailingsReport, ориентируясь по якорям."""
    if openpyxl is None:
        raise RuntimeError("openpyxl не установлен")
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[wb.sheetnames[0]]
    grid = [list(r) for r in ws.iter_rows(values_only=True)]
    n = len(grid)

    rep = TailingsReport(source_name=path.split("/")[-1], plant=plant)

    def cell(r, c):
        if 0 <= r < n and 0 <= c < len(grid[r]):
            return grid[r][c]
        return None

    def label_col(r):
        """Находит колонку с текстовой подписью (обычно 1, но бывает сдвиг)."""
        for c in range(0, 4):
            v = _norm(cell(r, c))
            if v:
                return c
        return 1

    # 1) Шихта руд (feed) и Отвальные/породные хвосты.
    for r in range(n):
        lc = label_col(r)
        lab = _norm(cell(r, lc)).lower()
        if lab.startswith("шихта руд") and rep.feed_mass_t == 0:
            rep.feed_mass_t = _num(cell(r, lc + 1))
        # первая встреченная строка хвостов с массой и содержанием
        if (lab.startswith("отвальные хвосты") or lab.startswith("хвосты породные")
                or lab.startswith("хвосты породны")):
            mass = _num(cell(r, lc + 1))
            e28p = _num(cell(r, lc + 2))
            e28t = _num(cell(r, lc + 3))
            e29p = _num(cell(r, lc + 4))
            e29t = _num(cell(r, lc + 5))
            # берём строку, где реально есть содержание (не просто масса)
            if e28t > 0 or e29t > 0:
                if rep.tail_mass_t == 0:
                    rep.tail_mass_t = mass
                    rep.tail_el28_pct, rep.tail_el28_t = e28p, e28t
                    rep.tail_el29_pct, rep.tail_el29_t = e29p, e29t

    # 2) Таблица распределения по классам крупности.
    #    Ищем шапку 'Класс крупности' и читаем следующие строки, пока не 'Итого'.
    gran_header = None
    for r in range(n):
        lc = label_col(r)
        if _norm(cell(r, lc)).lower().startswith("класс крупности"):
            gran_header = r
            break
    class_rows: dict = {}
    if gran_header is not None:
        for r in range(gran_header + 1, min(gran_header + 12, n)):
            lc = label_col(r)
            lab = _norm(cell(r, lc))
            if lab.lower().startswith("итого"):
                break
            size = None
            for s in SIZE_CLASSES:
                if _norm(lab).replace(" ", "") == s.replace(" ", ""):
                    size = s
                    break
            if size is None:
                continue
            scr = SizeClassRow(
                size=size,
                mass_share_pct=_num(cell(r, lc + 1)),
                el28_share_pct=_num(cell(r, lc + 2)),
                el28_t=_num(cell(r, lc + 3)),
                el29_share_pct=_num(cell(r, lc + 4)),
                el29_t=_num(cell(r, lc + 5)),
            )
            class_rows[size] = scr

    # 3) Минералогические блоки. Каждый блок начинается заголовком-классом
    #    (напр. '+71', '-71 + 45 мкм') в колонке подписи, ниже 'Доля потерь …'.
    for r in range(n):
        lc = label_col(r)
        lab = _norm(cell(r, lc))
        size = _canon_size(lab)
        if size is None:
            continue
        # подтверждаем, что это минералогический блок: рядом 'Доля потерь'
        header_txt = " ".join(_norm(cell(r, c)).lower() for c in range(lc, lc + 6))
        if "доля потерь" not in header_txt:
            continue
        m28 = MineralSplit()
        m29 = MineralSplit()
        for rr in range(r + 1, min(r + 14, n)):
            llc = label_col(rr)
            mlab = _norm(cell(rr, llc))
            if mlab.lower().startswith("итого"):
                # ниже могут быть 'Извлекаемый/Не извлекаемый металл'
                pass
            # колонки: llc=подпись, +2=доля%Ni, +3=т Ni, +4=доля%Cu, +5=т Cu
            if mlab.lower().startswith("извлекаемый металл"):
                m28.recoverable_t = _num(cell(rr, llc + 3))
                m29.recoverable_t = _num(cell(rr, llc + 5))
                continue
            if mlab.lower().startswith("не извлекаемый"):
                m28.unrecoverable_t = _num(cell(rr, llc + 3))
                m29.unrecoverable_t = _num(cell(rr, llc + 5))
                break
            canon = _canon_mineral(mlab)
            if canon is None:
                continue
            t28 = _num(cell(rr, llc + 3))
            t29 = _num(cell(rr, llc + 5))
            if t28:
                m28.minerals[canon] = m28.minerals.get(canon, 0.0) + t28
            if t29:
                m29.minerals[canon] = m29.minerals.get(canon, 0.0) + t29
        scr = class_rows.get(size) or SizeClassRow(size=size)
        scr.min28 = m28
        scr.min29 = m29
        class_rows[size] = scr

    rep.classes = [class_rows[s] for s in SIZE_CLASSES if s in class_rows]

    # 4) Валидация / предупреждения.
    if rep.tail_mass_t == 0:
        rep.warnings.append("Не найдена масса отвальных хвостов.")
    if not rep.classes:
        rep.warnings.append("Не найдено распределение по классам крупности.")
    if rep.tail_el28_t == 0 and rep.tail_el29_t == 0:
        rep.warnings.append("Нулевые потери металла — проверьте файл.")
    return rep


# --- docx / pdf ---------------------------------------------------------------

def read_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    out = []
    for p in d.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


if __name__ == "__main__":
    base = os.environ.get('CASE_DIR', 'data/examples')
    files = [
        ("Пример 1/Хвосты КГМК.xlsx", "КГМК"),
        ("Пример 2/Хвосты НОФ Вкр.xlsx", "НОФ-вкр"),
        ("Пример 3/Хвосты НОФ мед.xlsx", "НОФ-мед"),
        ("Пример 4/Хвосты ТОФ_2.xlsx", "ТОФ"),
    ]
    for fn, plant in files:
        rep = parse_tailings_xlsx(f"{base}/{fn}", plant)
        print(f"\n### {plant}: хвосты {rep.tail_mass_t:,.0f} т | "
              f"Ni {rep.tail_el28_t:,.0f} т ({rep.tail_el28_pct:.3f}%) | "
              f"Cu {rep.tail_el29_t:,.0f} т ({rep.tail_el29_pct:.3f}%)")
        for c in rep.classes:
            print(f"  {c.size:>9}: Ni {c.el28_t:7.1f}т "
                  f"(раскр {c.min28.liberated_t():6.1f} / закр {c.min28.locked_t():6.1f}) | "
                  f"Cu {c.el29_t:6.1f}т "
                  f"(раскр {c.min29.liberated_t():6.1f} / закр {c.min29.locked_t():6.1f})")
        if rep.warnings:
            print("  WARN:", rep.warnings)
