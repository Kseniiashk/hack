"""
Экспорт результатов: бизнес-отчёт docx, задачи csv/json, дорожная карта.
Форматы соответствуют ТЗ (docx/PDF-подобный отчёт, CSV/JSON для Jira/YouTrack).
"""
from __future__ import annotations
import os
import csv
import json
import io

from core.generate import Hypothesis
from core.diagnose import Diagnosis


def hypotheses_to_json(hyps: list) -> str:
    from dataclasses import asdict
    return json.dumps([asdict(h) for h in hyps], ensure_ascii=False, indent=2)


def hypotheses_to_csv(hyps: list) -> str:
    """CSV в формате задач (Jira/YouTrack): Summary, Description, Priority, …"""
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Rank", "Summary", "Score", "Ожид.эффект_M$", "Реализуемость",
                "Новизна", "Риск", "Класс_крупности", "Передел",
                "Обоснование", "Критерий_успеха", "Оборудование"])
    for i, h in enumerate(hyps, 1):
        w.writerow([i, h.title, f"{h.score:.2f}", f"{h.value_musd:.1f}",
                    f"{h.feasibility:.2f}", f"{h.novelty:.2f}", f"{h.risk:.2f}",
                    h.size_class, h.stage, h.rationale, h.success_criterion,
                    "; ".join(h.equipment)])
    return buf.getvalue()


def build_docx(diag: Diagnosis, hyps: list, out_path: str, top_n: int = None):
    """Формирует бизнес-отчёт .docx с ранжированными гипотезами и обоснованием."""
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    doc.add_heading("Фабрика гипотез — отчёт по снижению потерь металла с хвостами", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Объект: {diag.plant or diag.report_name}\n").bold = True
    p.add_run(f"Источник данных: {diag.report_name}")

    # --- Резюме диагностики ---
    doc.add_heading("1. Диагностика потерь", level=1)
    ceiling = next((f for f in diag.findings if f.code == "RECOVERABLE_CEILING"), None)
    if ceiling:
        doc.add_paragraph(ceiling.detail)
    doc.add_paragraph(
        f"Всего потенциально извлекаемо: {diag.total_recoverable_lost_t_ni:,.0f} т Ni и "
        f"{diag.total_recoverable_lost_t_cu:,.0f} т Cu. "
        f"Верхняя граница экономического эффекта ≈ {diag.total_value_musd:.1f} млн $/год.")

    doc.add_heading("Ключевые находки", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Причина", "Класс, мкм", "Эффект, M$/год", "Описание")
    for f in diag.findings:
        if f.code == "RECOVERABLE_CEILING":
            continue
        row = tbl.add_row().cells
        row[0].text = f.code
        row[1].text = f.size_class
        row[2].text = f"{f.value_musd:.1f}"
        row[3].text = f.headline

    # --- Ранжированные гипотезы ---
    doc.add_heading("2. Ранжированные гипотезы", level=1)
    show = hyps[:top_n] if top_n else hyps
    for i, h in enumerate(show, 1):
        doc.add_heading(f"{i}. {h.title}", level=2)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Score {h.score:.2f} | эффект ≈ {h.value_musd:.1f} млн $/год | "
            f"реализуемость {h.feasibility:.2f} | новизна {h.novelty:.2f} | "
            f"риск {h.risk:.2f} (тех {h.risk_tech:.2f}/экон {h.risk_econ:.2f}) | "
            f"класс {h.size_class} | передел: {h.stage}").italic = True

        doc.add_paragraph("Обоснование:", style="Intense Quote")
        llm_txt = getattr(h, "llm_rationale", "")
        doc.add_paragraph(llm_txt if llm_txt else h.rationale)

        if h.equipment:
            doc.add_paragraph(f"Оборудование: {', '.join(h.equipment)}")
        if h.success_criterion:
            doc.add_paragraph(f"Критерий успеха: {h.success_criterion}")
        if h.roadmap:
            doc.add_paragraph("Дорожная карта проверки:")
            for step in h.roadmap:
                doc.add_paragraph(step, style="List Number")
        if h.citations:
            doc.add_paragraph("Источники (литература):")
            for c in h.citations:
                pg = f", с.{c['page']}" if c.get("page") else ""
                doc.add_paragraph(f"«{c['snippet']}» — {c['source']}{pg}",
                                  style="List Bullet")

        # Протокол пилота для топ-5 гипотез.
        if i <= 5:
            from core.pilot import pilot_card
            pc = pilot_card(h)
            doc.add_paragraph("Протокол пилотной проверки (A/B):", style="Intense Quote")
            rows = [("Дизайн", pc["design"]), ("Длительность", pc["duration"])]
            if pc.get("capex"):
                rows.append(("Капзатраты", pc["capex"]))
            rows += [("Критерий успеха", pc["success"]), ("GO", pc["go"]),
                     ("STOP", pc["stop"]), ("Мерить ежедневно", pc["daily"])]
            for k, v in rows:
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Примечание: экономические оценки — ориентировочные (условные цены "
                 "Ni/Cu), служат для приоритизации. Элемент 28 = Ni, Элемент 29 = Cu.").italic = True

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    from core.ingest import parse_tailings_xlsx
    from core.diagnose import diagnose
    from core.generate import generate
    base = os.environ.get('CASE_DIR', 'data/examples')
    rep = parse_tailings_xlsx(f"{base}/Пример 1/Хвосты КГМК.xlsx", "КГМК")
    diag = diagnose(rep)
    hyps = generate(diag)
    out = os.path.join(os.path.dirname(__file__), "..", "exports", "КГМК_отчёт.docx")
    build_docx(diag, hyps, out, top_n=8)
    print("saved:", out)
    with open(out.replace(".docx", ".csv"), "w", encoding="utf-8") as f:
        f.write(hypotheses_to_csv(hyps))
    print("csv saved")
